"""
DOCX report builder for survey analysis.
Creates professional Word documents with tables and embedded charts.
"""

import logging
from pathlib import Path
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

logger = logging.getLogger(__name__)

# Colors
HEADER_BG = '2E5B88'
HEADER_TEXT = 'FFFFFF'
ALT_ROW_BG = 'F5F7FA'
BORDER_COLOR = 'D0D5DD'


def _set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:color="{BORDER_COLOR}"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:color="{BORDER_COLOR}"/>'
                          f'<w:left w:val="single" w:sz="4" w:color="{BORDER_COLOR}"/>'
                          f'<w:right w:val="single" w:sz="4" w:color="{BORDER_COLOR}"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


class ReportBuilder:
    """Builds a DOCX survey report incrementally."""
    
    def __init__(self, title: str = "Raport z badania ankietowego"):
        self.doc = Document()
        self._setup_styles()
        self._add_title_page(title)
    
    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x2D, 0x34, 0x36)
        
        # Heading styles
        for level, size, color in [(1, 16, '2E5B88'), (2, 13, '2E5B88'), (3, 11, '3D5A80')]:
            style_name = f'Heading {level}'
            if style_name in self.doc.styles:
                h_style = self.doc.styles[style_name]
                h_style.font.name = 'Calibri'
                h_style.font.size = Pt(size)
                h_style.font.bold = True
                h_style.font.color.rgb = RGBColor.from_string(color)
        
        # Set margins
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)
    
    def _add_title_page(self, title: str):
        """Add title page."""
        # Add spacing before title
        for _ in range(6):
            self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string('2E5B88')
        
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Wyniki automatycznej analizy")
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor.from_string('636E72')
        
        self.doc.add_page_break()
    
    def add_section(self, title: str, level: int = 1):
        """Add a section heading."""
        self.doc.add_heading(title, level=level)
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a text paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p
    
    def add_table(self, df: pd.DataFrame, title: str = "", 
                  max_col_width: Optional[float] = None):
        """Add a formatted table from DataFrame."""
        if title:
            p = self.doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string('2D3436')
        
        n_rows = len(df)
        n_cols = len(df.columns)
        
        table = self.doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        for j, col_name in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = str(col_name)
            _set_cell_shading(cell, HEADER_BG)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'
        
        # Data rows
        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.cell(i + 1, j)
                val = df.iloc[i, j]
                cell.text = str(val) if not pd.isna(val) else ''
                
                # Alternating row colors
                if i % 2 == 1:
                    _set_cell_shading(cell, ALT_ROW_BG)
                
                for paragraph in cell.paragraphs:
                    if j > 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Calibri'
        
        self.doc.add_paragraph()  # spacing
        return table
    
    def add_chart(self, chart_buf: BytesIO, width: float = 6.0):
        """Add a chart image from BytesIO buffer."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_buf, width=Inches(width))
        self.doc.add_paragraph()  # spacing
    
    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()
    
    def save(self, filepath: str):
        """Save the document."""
        self.doc.save(filepath)
        logger.info(f"Report saved to {filepath}")
